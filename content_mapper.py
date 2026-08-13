#!/usr/bin/env python3
"""
Moodle Course Content Mapper

Consumes the outputs of cloudpedagogy-moodle-course-auditor:

    input/<course>/
        audit/
            sections.csv
            activities.csv
            content_placement_inventory.csv
            course_summary.csv          (optional)
        extracted_files/
            ... recovered Moodle resources ...

Produces browsable HTML and editable Word content maps plus CSV/report outputs.

Design goals
------------
* Preserve Moodle section and course-item order.
* Make as few interpretive decisions as possible.
* Use module_id / section metadata from the audit as the primary mapping key.
* Link Moodle-hosted resources to the actual extracted files.
* Preserve external URLs (Panopto, SharePoint, Forms, etc.).
* Flag unresolved resources instead of inventing mappings.
* Use only python-docx beyond the Python standard library (for Word output).

Example
-------
    python3 content_mapper.py input/literature-review

Portable bundle:
    python3 content_mapper.py input/literature-review --bundle
"""

from __future__ import annotations

import argparse
import csv
import html
import os
import re
import shutil
import sys
import unicodedata
from datetime import datetime, timezone
from difflib import SequenceMatcher
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple
from urllib.parse import quote

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError:
    Document = None


VERSION = "0.5.0"


def clean(value) -> str:
    if value is None:
        return ""
    value = str(value).strip()
    if value.lower() in {"nan", "none", "null"}:
        return ""
    return value


def intish(value, default: int = 0) -> int:
    try:
        return int(float(clean(value)))
    except (ValueError, TypeError):
        return default


def truthy(value) -> bool:
    return clean(value).lower() in {"1", "true", "yes", "y"}


def read_csv(path: Path) -> List[dict]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as fh:
        return list(csv.DictReader(fh))


def safe_filename(name: str, fallback: str = "resource") -> str:
    name = clean(name) or fallback
    name = name.replace("/", "_").replace("\\", "_")
    name = re.sub(r"[\x00-\x1f\x7f]", "", name)
    return name[:220] or fallback


def slugify(text: str) -> str:
    text = clean(text).lower()
    text = re.sub(r"[^a-z0-9]+", "-", text).strip("-")
    return text or "section"


def html_path(path: str) -> str:
    path = path.replace(os.sep, "/")
    return "/".join(part if part in {"", ".", ".."} else quote(part) for part in path.split("/"))


def first_present(row: dict, *keys: str) -> str:
    for key in keys:
        value = clean(row.get(key))
        if value:
            return value
    return ""


def normalize_filename_key(name: str) -> str:
    """
    Conservative filename normalization used only for matching.
    """
    name = clean(name)
    if not name:
        return ""

    name = unicodedata.normalize("NFKC", name)
    translations = {
        "\u00a0": " ",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u2212": "-",
    }
    for old, new in translations.items():
        name = name.replace(old, new)

    name = re.sub(r"\s+", " ", name).strip()
    name = re.sub(r"\s*([._(),-])\s*", r"\1", name)
    return name.casefold()


def fuzzy_filename_score(a: str, b: str) -> float:
    """Similarity score for final-resort filename matching."""
    return SequenceMatcher(
        None,
        normalize_filename_key(a),
        normalize_filename_key(b),
    ).ratio()


@dataclass
class LinkResult:
    href: str = ""
    status: str = ""
    source_path: str = ""
    bundled_path: str = ""


@dataclass
class MapRow:
    section_number: int
    section_name: str
    activity_order: int
    module_id: str
    activity_type: str
    activity_name: str
    visible: str
    content_category: str
    content_subtype: str
    provider: str
    filename_or_title: str
    link_type: str
    href: str
    mapping_status: str
    classification_confidence: str
    association_status: str
    modified_year: str = ""
    duplicate_candidate: str = ""


REQUIRED_AUDIT_FILES = ("sections.csv", "activities.csv", "content_placement_inventory.csv")


def validate_input(course_dir: Path) -> Tuple[Path, Path]:
    audit_dir = course_dir / "audit"
    extracted_dir = course_dir / "extracted_files"
    if not course_dir.exists():
        raise FileNotFoundError(f"Input course directory not found: {course_dir}")
    if not audit_dir.is_dir():
        raise FileNotFoundError(f"Expected audit directory not found: {audit_dir}")
    missing = [name for name in REQUIRED_AUDIT_FILES if not (audit_dir / name).exists()]
    if missing:
        raise FileNotFoundError("Required audit file(s) missing: " + ", ".join(missing))
    if not extracted_dir.is_dir():
        raise FileNotFoundError(f"Expected extracted_files directory not found: {extracted_dir}")
    return audit_dir, extracted_dir


def discover_course_title(audit_dir: Path, fallback: str) -> str:
    rows = read_csv(audit_dir / "course_summary.csv")
    if rows:
        row = rows[0]
        for key in ("course_fullname_from_xml", "fullname", "course_fullname", "full_name", "course_name", "name"):
            value = clean(row.get(key))
            if value:
                return value
        for row in rows:
            key = first_present(row, "field", "metric", "key", "name").lower()
            value = first_present(row, "value", "result")
            if key in {"fullname", "full name", "course name"} and value:
                return value
    return fallback.replace("-", " ").replace("_", " ").title()



def extract_modified_year(*rows: dict) -> str:
    """Return an item-level modification year when explicitly available."""
    keys = (
        "modification_year", "modified_year", "last_modified_year",
        "timemodified", "time_modified", "modified", "last_modified",
        "modified_at", "date_modified", "updated_at",
    )
    current_year = datetime.now(timezone.utc).year
    for row in rows:
        if not row:
            continue
        for key in keys:
            raw = clean(row.get(key))
            if not raw:
                continue
            match = re.search(r"\b(?:19|20)\d{2}\b", raw)
            if match:
                year = int(match.group(0))
                if 1990 <= year <= current_year + 1:
                    return str(year)
            if re.fullmatch(r"\d{9,13}", raw):
                try:
                    stamp = int(raw)
                    if stamp > 10_000_000_000:
                        stamp //= 1000
                    year = datetime.fromtimestamp(stamp, tz=timezone.utc).year
                    if 1990 <= year <= current_year + 1:
                        return str(year)
                except (ValueError, OSError, OverflowError):
                    pass
    return ""


def load_duplicate_candidates(audit_dir: Path) -> set[str]:
    """Read optional duplicate_activity_inventory.csv using explicit IDs only."""
    path = audit_dir / "duplicate_activity_inventory.csv"
    if not path.exists():
        return set()
    result: set[str] = set()
    for row in read_csv(path):
        module_id = first_present(row, "module_id", "course_module_id", "cmid", "activity_id")
        if not module_id:
            continue
        flag = first_present(
            row, "duplicate_candidate", "is_duplicate", "duplicate", "is_duplicate_candidate"
        )
        if flag and flag.lower() in {"0", "false", "no", "n"}:
            continue
        result.add(module_id)
    return result


def file_extension_label(name: str) -> str:
    suffix = Path(clean(name)).suffix.lower().lstrip(".")
    return suffix.upper() if suffix else ""


def unique_sorted(values) -> List[str]:
    return sorted({clean(v) for v in values if clean(v)}, key=lambda value: value.casefold())


def item_visibility_label(value: str) -> str:
    value = clean(value)
    if not value:
        return "Unknown"
    return "Visible" if truthy(value) else "Hidden"


def mapping_qa_label(status: str, has_href: bool) -> str:
    status = clean(status)
    if status.startswith("ambiguous-"):
        return "Ambiguous"
    if status in {"unresolved-file", "no-filename"}:
        return "Unresolved"
    if has_href:
        return "Linked"
    if status in {"activity-only", "no-linkable-resource"}:
        return "No separate resource"
    return "Other"


def real_course_item_count(section: dict) -> int:
    """Count displayed course items, excluding Moodle labels used as subheadings."""
    return sum(
        1 for item in section.get("activities", [])
        if clean(item.get("activity_type")).lower() != "label"
    )


class ResourceIndex:
    def __init__(self, extracted_dir: Path):
        self.by_basename: Dict[str, List[Path]] = defaultdict(list)
        self.by_basename_lower: Dict[str, List[Path]] = defaultdict(list)
        self.by_normalized: Dict[str, List[Path]] = defaultdict(list)
        self.all_files: List[Path] = []

        ignored_names = {"extraction_report.md", "resource_manifest.csv"}

        for path in extracted_dir.rglob("*"):
            if not path.is_file() or path.name in ignored_names:
                continue

            rp = path.resolve()
            self.all_files.append(rp)
            self.by_basename[path.name].append(rp)
            self.by_basename_lower[path.name.casefold()].append(rp)

            norm = normalize_filename_key(path.name)
            if norm:
                self.by_normalized[norm].append(rp)

        self._sort_candidates()

    def _sort_candidates(self) -> None:
        def preference(path: Path):
            text = path.as_posix()
            if "/files_by_moodle_context/" in text:
                tier = 0
            elif "/resource_bundle/" in text:
                tier = 1
            elif "/resource_bundle_by_type/" in text:
                tier = 2
            else:
                tier = 3
            return (tier, len(path.parts), text.casefold())

        for mapping in (
            self.by_basename,
            self.by_basename_lower,
            self.by_normalized,
        ):
            for key in mapping:
                mapping[key] = sorted(mapping[key], key=preference)

        self.all_files = sorted(self.all_files, key=preference)

    @staticmethod
    def _collapse_duplicate_views(paths: List[Path]) -> List[Path]:
        """
        Collapse duplicate extraction views of the same filename while
        preserving the preferred files_by_moodle_context copy.
        """
        seen = set()
        out = []
        for path in paths:
            key = path.name.casefold()
            if key in seen:
                continue
            seen.add(key)
            out.append(path)
        return out

    def match(self, filename: str) -> Tuple[Optional[Path], str]:
        """
        Accuracy-first matching order:
        1. Exact basename
        2. Case-insensitive exact basename
        3. Conservative normalized basename
        4. Very-high-confidence fuzzy basename, same extension only

        Ambiguous matches are never auto-linked.
        """
        filename = clean(filename)
        if not filename:
            return None, "no-filename"

        exact = self.by_basename.get(filename, [])
        if exact:
            return exact[0], "exact-filename"

        ci = self.by_basename_lower.get(filename.casefold(), [])
        if ci:
            return ci[0], "case-insensitive-filename"

        norm_key = normalize_filename_key(filename)
        normalized = self._collapse_duplicate_views(
            self.by_normalized.get(norm_key, [])
        )
        if len(normalized) == 1:
            return normalized[0], "normalized-filename"
        if len(normalized) > 1:
            return None, "ambiguous-normalized-filename"

        target_ext = Path(filename).suffix.casefold()
        candidates = []
        for path in self.all_files:
            if path.suffix.casefold() != target_ext:
                continue
            score = fuzzy_filename_score(filename, path.name)
            if score >= 0.985:
                candidates.append((score, path))

        candidates.sort(key=lambda x: (-x[0], x[1].as_posix().casefold()))
        if not candidates:
            return None, "unresolved-file"

        best_score, best_path = candidates[0]
        runner_up = candidates[1][0] if len(candidates) > 1 else 0.0

        if len(candidates) > 1 and (best_score - runner_up) < 0.01:
            distinct = {p.name.casefold() for _, p in candidates[:3]}
            if len(distinct) > 1:
                return None, "ambiguous-fuzzy-filename"

        return best_path, f"high-confidence-fuzzy:{best_score:.3f}"


def build_activity_lookup(activities: Sequence[dict]) -> Dict[str, dict]:
    return {clean(r.get("module_id")): r for r in activities if clean(r.get("module_id"))}


def build_placements_lookup(placements: Sequence[dict]) -> Dict[str, List[dict]]:
    result = defaultdict(list)
    for row in placements:
        mid = clean(row.get("module_id"))
        if mid:
            result[mid].append(row)
    return result


def section_activity_ids(section: dict, activities: Sequence[dict]) -> List[str]:
    seq = [x.strip() for x in clean(section.get("activity_sequence")).split(",") if x.strip()]
    if seq:
        return seq
    sid = clean(section.get("section_id"))
    sn = clean(section.get("section_number"))
    return [clean(r.get("module_id")) for r in activities if clean(r.get("module_id")) and ((sid and clean(r.get("section_id")) == sid) or (sn and clean(r.get("section_number")) == sn))]


def placement_display_name(row: dict, activity: dict) -> str:
    return first_present(row, "filename_or_title", "activity_name") or first_present(activity, "activity_name") or "Resource"


def make_link(row: dict, resource_index: ResourceIndex, output_dir: Path, bundle: bool, bundle_dir: Path, registry: Dict[str, int]) -> LinkResult:
    url = first_present(row, "canonical_url", "url", "url_or_reference")
    if url:
        return LinkResult(href=url, status="external-url")
    filename = first_present(row, "filename_or_title")
    if not filename:
        return LinkResult(status="no-linkable-resource")
    source, status = resource_index.match(filename)
    if not source:
        return LinkResult(status=status)
    if bundle:
        bundle_dir.mkdir(parents=True, exist_ok=True)
        base = safe_filename(source.name)
        key = base.lower()
        registry[key] += 1
        n = registry[key]
        target_name = base if n == 1 else f"{Path(base).stem}__{n}{Path(base).suffix}"
        target = bundle_dir / target_name
        shutil.copy2(source, target)
        return LinkResult(href=html_path(os.path.relpath(target, output_dir)), status=f"{status};bundled", source_path=str(source), bundled_path=str(target))
    return LinkResult(href=html_path(os.path.relpath(source, output_dir)), status=status, source_path=str(source))


def map_course(
    sections,
    activities,
    placements,
    resource_index,
    output_dir,
    bundle,
    include_hidden,
    duplicate_candidates=None,
):
    activity_lookup = build_activity_lookup(activities)
    placements_lookup = build_placements_lookup(placements)
    duplicate_candidates = duplicate_candidates or set()

    section_models, flat_rows, unresolved = [], [], []
    bundle_dir = output_dir / "resources"
    registry = defaultdict(int)

    for section in sorted(
        sections,
        key=lambda row: (
            intish(row.get("section_number"), 999999),
            clean(row.get("section_name")),
        ),
    ):
        sn = intish(section.get("section_number"), 0)
        sname = clean(section.get("section_name")) or f"Section {sn}"
        if not include_hidden and not truthy(section.get("visible")):
            continue

        mapped = []
        for order_index, mid in enumerate(section_activity_ids(section, activities), 1):
            activity = activity_lookup.get(mid)
            if not activity:
                unresolved.append({
                    "section_number": sn,
                    "section_name": sname,
                    "module_id": mid,
                    "activity_name": "",
                    "reason": "module_id in section sequence but absent from activities.csv",
                })
                continue

            if not include_hidden and not truthy(activity.get("visible")):
                continue

            atype = clean(activity.get("activity_type")) or "activity"
            aname = clean(activity.get("activity_name")) or f"{atype} {mid}"
            visible_value = clean(activity.get("visible"))
            placement_rows = placements_lookup.get(mid, [])
            item_links = []
            modified_year = extract_modified_year(activity, *placement_rows)
            duplicate_candidate = "Yes" if mid in duplicate_candidates else ""

            for placement in placement_rows:
                link = make_link(
                    placement, resource_index, output_dir, bundle, bundle_dir, registry
                )
                display = placement_display_name(placement, activity)
                provider = first_present(placement, "provider", "hosting_type")
                category = clean(placement.get("content_category"))
                subtype = clean(placement.get("content_subtype"))
                confidence = clean(placement.get("classification_confidence"))
                association = clean(placement.get("association_status"))
                external_source = first_present(
                    placement, "canonical_url", "url", "url_or_reference"
                )
                link_type = (
                    "external" if external_source and link.href
                    else ("local-file" if link.href else "none")
                )

                item_links.append({
                    "display_name": display,
                    "href": link.href,
                    "link_type": link_type,
                    "mapping_status": link.status,
                    "provider": provider,
                    "content_category": category,
                    "content_subtype": subtype,
                    "file_extension": file_extension_label(display),
                })

                flat_rows.append(MapRow(
                    sn, sname, order_index, mid, atype, aname, visible_value,
                    category, subtype, provider, display, link_type, link.href,
                    link.status, confidence, association, modified_year,
                    duplicate_candidate
                ))

                if not link.href and link.status in {
                    "unresolved-file", "no-filename",
                    "ambiguous-normalized-filename", "ambiguous-fuzzy-filename"
                }:
                    unresolved.append({
                        "section_number": sn,
                        "section_name": sname,
                        "module_id": mid,
                        "activity_name": aname,
                        "reason": f"{link.status}: {display}",
                    })

            if not item_links:
                fallback_url = first_present(activity, "xml_external_links_sample")
                if fallback_url:
                    provider = clean(activity.get("xml_external_domains"))
                    item_links.append({
                        "display_name": aname,
                        "href": fallback_url,
                        "link_type": "external",
                        "mapping_status": "activity-xml-url-fallback",
                        "provider": provider,
                        "content_category": "",
                        "content_subtype": "",
                        "file_extension": "",
                    })
                    flat_rows.append(MapRow(
                        sn, sname, order_index, mid, atype, aname, visible_value,
                        "", "", provider, aname, "external", fallback_url,
                        "activity-xml-url-fallback", "", "", modified_year,
                        duplicate_candidate
                    ))
                else:
                    flat_rows.append(MapRow(
                        sn, sname, order_index, mid, atype, aname, visible_value,
                        "", "", "", "", "none", "", "activity-only", "", "",
                        modified_year, duplicate_candidate
                    ))

            mapped.append({
                "module_id": mid,
                "activity_type": atype,
                "activity_name": aname,
                "order": order_index,
                "items": item_links,
                "family": course_item_family(atype),
                "type_label": type_label(atype),
                "visible": visible_value,
                "visibility_label": item_visibility_label(visible_value),
                "providers": unique_sorted(i.get("provider") for i in item_links),
                "link_types": unique_sorted(i.get("link_type") for i in item_links),
                "file_extensions": unique_sorted(i.get("file_extension") for i in item_links),
                "qa_values": unique_sorted(
                    mapping_qa_label(i.get("mapping_status"), bool(i.get("href")))
                    for i in item_links
                ),
                "modified_year": modified_year,
                "duplicate_candidate": duplicate_candidate,
            })

        section_models.append({
            "section_number": sn,
            "section_name": sname,
            "summary": clean(section.get("section_summary_text_from_xml")),
            "activities": mapped,
        })
    return section_models, flat_rows, unresolved


def add_docx_hyperlink(paragraph, text: str, target: str):
    """Add a clickable local/external hyperlink to a python-docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def write_docx(path: Path, title: str, sections, rows) -> None:
    """Create an editable Word working copy of the audited Moodle course-item structure."""
    if Document is None:
        raise RuntimeError(
            "Word output requires python-docx. Install it with: "
            "python3 -m pip install python-docx"
        )

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 2"].font.name = "Arial"

    doc.add_heading(title, 0)
    subtitle = doc.add_paragraph()
    r = subtitle.add_run("Editable Moodle Course Content Map")
    r.bold = True

    note = doc.add_paragraph()
    note.add_run(
        "Purpose: an editable representation of the existing Moodle course structure. "
        "The mapper has not pedagogically reclassified or reorganised the content. "
        "Move or rename headings, course items and resource entries in this document as part of course redesign; "
        "the underlying extracted resources are not changed."
    )

    link_note = doc.add_paragraph()
    link_note.add_run("Links: ").bold = True
    link_note.add_run(
        "local resource hyperlinks point to the extracted files used by the HTML map. "
        "Keep the surrounding course-map folders in their relative positions, or generate a portable "
        "bundle with --bundle if the package needs to be moved elsewhere."
    )

    terminology_note = doc.add_paragraph()
    terminology_note.add_run("Terminology: ").bold = True
    terminology_note.add_run(
        "'Course item' is used as the umbrella term for elements represented in the Moodle course. "
        "Where possible, the mapper also indicates whether an item is resource/content "
        "(for example a file, URL, page or book) or a learning activity "
        "(for example a forum, quiz or assignment)."
    )

    doc.add_heading("Contents / working outline", level=1)
    for s in sections:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f'{s["section_number"]}. {s["section_name"]}')

    doc.add_page_break()

    for s in sections:
        doc.add_heading(f'{s["section_number"]}. {s["section_name"]}', level=1)

        if s.get("summary"):
            p = doc.add_paragraph()
            lead = p.add_run("Current Moodle section description: ")
            lead.bold = True
            p.add_run(s["summary"])

        for activity in s["activities"]:
            atype = activity["activity_type"]
            aname = activity["activity_name"]

            if atype == "label":
                doc.add_heading(aname, level=2)
                continue

            p = doc.add_paragraph()
            badge = p.add_run(f"[{type_label(atype)}] ")
            badge.bold = True
            family = course_item_family(atype)
            family_run = p.add_run(f"{family}: ")
            family_run.italic = True
            name_run = p.add_run(aname)
            name_run.bold = True

            if activity["items"]:
                for item in activity["items"]:
                    rp = doc.add_paragraph(style="List Bullet 2")
                    display = item["display_name"]
                    if item["href"]:
                        add_docx_hyperlink(rp, display, item["href"])
                    else:
                        rp.add_run(display)
                        unresolved = rp.add_run("  [Resource not resolved]")
                        unresolved.italic = True

                    meta = " | ".join(
                        x for x in (
                            item.get("content_subtype") or item.get("content_category"),
                            item.get("provider"),
                        ) if x
                    )
                    if meta:
                        m = rp.add_run(f"  ({meta})")
                        m.italic = True
            else:
                rp = doc.add_paragraph(style="List Bullet 2")
                rp.add_run("No separate file or URL was identified for this Moodle course item.").italic = True

    doc.add_page_break()
    doc.add_heading("Mapping notes", level=1)
    notes = [
        "Section order is taken from sections.csv.",
        "Course-item order is taken from each Moodle section's activity sequence.",
        "Existing Moodle labels are retained as subheadings.",
        "Local resources are linked only when the mapper resolves a sufficiently reliable file match.",
        "External URLs are retained from the audit data.",
        "No semantic or pedagogic remapping is performed in this document.",
    ]
    for n in notes:
        doc.add_paragraph(n, style="List Bullet")

    doc.save(path)


def write_csv_output(path: Path, rows: Sequence[MapRow]) -> None:
    fields = list(MapRow.__dataclass_fields__.keys())
    with path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({f: getattr(row, f) for f in fields})


def write_unresolved_csv(path: Path, rows: Sequence[dict]) -> None:
    fields = ["section_number", "section_name", "module_id", "activity_name", "reason"]
    with path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({f: row.get(f, "") for f in fields})


def write_report(path: Path, title: str, sections, rows, unresolved, bundle: bool) -> None:
    local = sum(r.link_type == "local-file" for r in rows)
    external = sum(r.link_type == "external" for r in rows)
    total_course_items = sum(real_course_item_count(s) for s in sections)
    exact = sum(r.mapping_status == "exact-filename" for r in rows)
    ci = sum(r.mapping_status == "case-insensitive-filename" for r in rows)
    normalized = sum(r.mapping_status == "normalized-filename" for r in rows)
    fuzzy = sum(r.mapping_status.startswith("high-confidence-fuzzy:") for r in rows)
    ambiguous = sum(r.mapping_status.startswith("ambiguous-") for r in rows)

    text = [
        f"# Course Content Mapping Report - {title}", "",
        "## Summary", "",
        f"- Sections mapped: {len(sections)}",
        f"- Moodle course items represented: {total_course_items}",
        f"- Mapping rows: {len(rows)}",
        f"- Local extracted-resource links: {local}",
        f"- External links: {external}",
        f"- Unresolved mapping issues: {len(unresolved)}",
        f"- Portable resource bundle: {'Yes' if bundle else 'No'}", "",
        "## Resource-link QA", "",
        f"- Exact filename matches: {exact}",
        f"- Case-insensitive exact matches: {ci}",
        f"- Normalized filename matches: {normalized}",
        f"- High-confidence fuzzy matches: {fuzzy}",
        f"- Ambiguous matches left unresolved: {ambiguous}", "",
        "## Mapping approach", "",
        "- Section order comes from `sections.csv`.",
        "- Course-item order comes from each section's `activity_sequence`.",
        "- Resources and URLs are associated by `module_id` from `content_placement_inventory.csv`.",
        "- Extracted files are resolved in this order: exact filename, case-insensitive exact filename, conservative normalized filename, then very-high-confidence same-extension fuzzy matching.",
        "- Ambiguous normalized or fuzzy matches are deliberately left unresolved.",
        "- Existing Moodle labels are used as subheadings and are not counted as displayed course items.",
        "- Optional HTML metadata is shown only when supported by item-level audit data.",
        "- No semantic or pedagogic remapping is attempted.",
        "- Unresolved resources are reported rather than guessed.", ""
    ]
    if unresolved:
        text += ["## Unresolved items", ""]
        for item in unresolved:
            text.append(f"- Section {item.get('section_number','')} - {item.get('activity_name') or item.get('module_id','')}: {item.get('reason','')}")
        text.append("")
    path.write_text("\n".join(text), encoding="utf-8")


def type_label(activity_type: str) -> str:
    """
    User-facing Moodle course-item type.

    Internally the auditor uses activity_type/activity_name for all course
    modules. The mapper presents these as "course items" because Moodle
    courses contain both learning activities and resources.
    """
    labels = {
        "resource": "File",
        "url": "Link",
        "label": "Heading",
        "forum": "Forum",
        "lti": "LTI",
        "coursework": "Assessment",
        "assign": "Assignment",
        "assignment": "Assignment",
        "folder": "Folder",
        "page": "Page",
        "book": "Book",
        "quiz": "Quiz",
        "lesson": "Lesson",
        "scorm": "SCORM",
        "h5pactivity": "H5P",
        "choice": "Choice",
        "feedback": "Feedback",
        "glossary": "Glossary",
        "wiki": "Wiki",
        "workshop": "Workshop",
        "chat": "Chat",
        "database": "Database",
    }
    return labels.get(activity_type.lower(), activity_type.title() or "Course item")


def course_item_family(activity_type: str) -> str:
    """Broad user-facing classification without altering the audit data."""
    atype = clean(activity_type).lower()

    resource_types = {
        "resource", "url", "folder", "page", "book", "label",
    }
    activity_types = {
        "forum", "quiz", "assign", "assignment", "coursework", "lesson",
        "scorm", "h5pactivity", "choice", "feedback", "glossary", "wiki",
        "workshop", "chat", "database", "lti",
    }

    if atype in resource_types:
        return "Resource/content"
    if atype in activity_types:
        return "Activity"
    return "Course item"


def render_html(title: str, sections, rows) -> str:
    """Render a portable review map with optional, data-driven filters."""
    toc, body = [], []

    all_items = [
        activity
        for section in sections
        for activity in section["activities"]
        if clean(activity.get("activity_type")).lower() != "label"
    ]

    families = unique_sorted(item.get("family") for item in all_items)
    types = unique_sorted(item.get("type_label") for item in all_items)
    providers = unique_sorted(
        provider for item in all_items for provider in item.get("providers", [])
    )
    link_types = unique_sorted(
        link_type for item in all_items
        for link_type in item.get("link_types", [])
        if link_type and link_type != "none"
    )
    extensions = unique_sorted(
        ext for item in all_items for ext in item.get("file_extensions", [])
    )
    years = sorted(
        {item.get("modified_year") for item in all_items if item.get("modified_year")},
        reverse=True,
    )
    qa_values = unique_sorted(
        value for item in all_items for value in item.get("qa_values", [])
    )
    visibility_values = unique_sorted(
        item.get("visibility_label")
        for item in all_items
        if item.get("visibility_label") not in {"", "Unknown"}
    )
    has_duplicates = any(
        item.get("duplicate_candidate") == "Yes" for item in all_items
    )

    def select_filter(filter_id: str, label: str, values: Sequence[str]) -> str:
        values = [clean(v) for v in values if clean(v)]
        if len(values) < 2:
            return ""
        options = ['<option value="">All</option>']
        options.extend(
            f'<option value="{html.escape(v, quote=True)}">{html.escape(v)}</option>'
            for v in values
        )
        return (
            f'<label class="filter-field" for="{html.escape(filter_id)}">'
            f'<span>{html.escape(label)}</span>'
            f'<select id="{html.escape(filter_id)}" data-filter="{html.escape(filter_id)}">'
            f'{"".join(options)}</select></label>'
        )

    controls = [
        select_filter("family", "Category", families),
        select_filter("type", "Item type", types),
        select_filter("provider", "Provider", providers),
        select_filter("linktype", "Link location", link_types),
        select_filter("extension", "File type", extensions),
        select_filter("year", "Modified year", years),
        select_filter("qa", "Link QA", qa_values),
        select_filter("visibility", "Visibility", visibility_values),
    ]
    controls = [c for c in controls if c]

    duplicate_control = (
        '<label class="check-field"><input id="duplicates-only" type="checkbox"> '
        'Duplicate candidates only</label>'
        if has_duplicates else ""
    )

    for section in sections:
        sid = f"section-{section['section_number']}-{slugify(section['section_name'])}"
        toc.append(
            f'<li><a href="#{html.escape(sid)}">'
            f'{html.escape(str(section["section_number"]))}. '
            f'{html.escape(section["section_name"])}</a></li>'
        )

        bits, subgroup_open = [], False
        for activity in section["activities"]:
            atype = activity["activity_type"]
            aname = activity["activity_name"]

            if atype == "label":
                if subgroup_open:
                    bits.append("</div>")
                bits.append(f'<div class="subgroup"><h3>{html.escape(aname)}</h3>')
                subgroup_open = True
                continue

            rendered_items = []
            for item in activity["items"]:
                display = html.escape(item["display_name"])
                meta = " · ".join(
                    clean(x) for x in (
                        item.get("content_subtype") or item.get("content_category"),
                        item.get("provider"),
                    ) if clean(x)
                )
                meta_html = (
                    f'<span class="item-meta">{html.escape(meta)}</span>' if meta else ""
                )

                if item["href"]:
                    target = (
                        ' target="_blank" rel="noopener noreferrer"'
                        if item["link_type"] == "external" else ""
                    )
                    rendered_items.append(
                        f'<li><a class="resource-link" '
                        f'href="{html.escape(item["href"], quote=True)}"{target}>'
                        f'{display}</a>{meta_html}</li>'
                    )
                else:
                    rendered_items.append(
                        f'<li><span>{display}</span>'
                        f'<span class="unresolved">Resource not resolved</span>'
                        f'{meta_html}</li>'
                    )

            items_block = (
                f'<ul class="resource-list">{"".join(rendered_items)}</ul>'
                if rendered_items else
                '<div class="activity-note">No separate file or URL was identified for this Moodle course item.</div>'
            )

            details = []
            if activity.get("modified_year"):
                details.append(("Last modified", activity["modified_year"]))
            if activity.get("providers"):
                details.append(("Provider", ", ".join(activity["providers"])))
            locations = [
                {"local-file": "Bundled/local", "external": "External"}.get(v, v)
                for v in activity.get("link_types", []) if v != "none"
            ]
            if locations:
                details.append(("Link location", ", ".join(locations)))
            if activity.get("file_extensions"):
                details.append(("File type", ", ".join(activity["file_extensions"])))
            if activity.get("visibility_label") not in {"", "Unknown"}:
                details.append(("Visibility", activity["visibility_label"]))
            if activity.get("duplicate_candidate") == "Yes":
                details.append(("Duplicate candidate", "Yes"))
            if activity.get("qa_values"):
                details.append(("Link QA", ", ".join(activity["qa_values"])))

            details_html = ""
            if details:
                rows_html = "".join(
                    f'<div><dt>{html.escape(label)}</dt><dd>{html.escape(value)}</dd></div>'
                    for label, value in details
                )
                details_html = (
                    '<details class="metadata"><summary>Review metadata</summary>'
                    f'<dl>{rows_html}</dl></details>'
                )

            search_text = " ".join(
                [aname, atype, activity.get("family", ""), activity.get("type_label", "")]
                + [i["display_name"] for i in activity["items"]]
                + [i["provider"] for i in activity["items"]]
                + activity.get("file_extensions", [])
                + activity.get("qa_values", [])
                + ([activity["modified_year"]] if activity.get("modified_year") else [])
            ).lower()

            attrs = {
                "data-search": search_text,
                "data-family": activity.get("family", ""),
                "data-type": activity.get("type_label", ""),
                "data-provider": "|".join(activity.get("providers", [])),
                "data-linktype": "|".join(activity.get("link_types", [])),
                "data-extension": "|".join(activity.get("file_extensions", [])),
                "data-year": activity.get("modified_year", ""),
                "data-qa": "|".join(activity.get("qa_values", [])),
                "data-visibility": activity.get("visibility_label", ""),
                "data-duplicate": "yes" if activity.get("duplicate_candidate") == "Yes" else "no",
            }
            attr_html = " ".join(
                f'{key}="{html.escape(clean(value), quote=True)}"'
                for key, value in attrs.items()
            )

            bits.append(
                f'<article class="activity" {attr_html}>'
                f'<div class="activity-title-row">'
                f'<span class="type-badge">{html.escape(activity.get("type_label") or type_label(atype))}</span>'
                f'<div class="item-heading">'
                f'<span class="family-label">{html.escape(activity.get("family") or course_item_family(atype))}</span>'
                f'<h4>{html.escape(aname)}</h4></div></div>'
                f'{items_block}{details_html}</article>'
            )

        if subgroup_open:
            bits.append("</div>")

        summary = section["summary"]
        summary_block = ""
        if summary:
            excerpt = summary if len(summary) <= 700 else summary[:697].rstrip() + "..."
            summary_block = (
                '<details class="section-summary"><summary>Current Moodle section description</summary>'
                f'<p>{html.escape(excerpt)}</p></details>'
            )

        displayed_count = real_course_item_count(section)
        body.append(
            f'<section class="course-section" id="{html.escape(sid)}">'
            f'<div class="section-heading"><div class="section-number">'
            f'{html.escape(str(section["section_number"]))}</div><div>'
            f'<h2>{html.escape(section["section_name"])}</h2>'
            f'<div class="section-count">{displayed_count} Moodle course items</div>'
            f'</div></div>{summary_block}'
            f'<div class="activity-list">{"".join(bits)}</div></section>'
        )

    total_course_items = sum(real_course_item_count(section) for section in sections)
    linked = sum(bool(row.href) for row in rows)
    external = sum(row.link_type == "external" for row in rows)

    filters_html = ""
    if controls or duplicate_control:
        filters_html = (
            '<div class="filters"><div class="filters-title">Review filters</div>'
            + "".join(controls)
            + duplicate_control
            + '<button id="clear-filters" type="button">Clear filters</button></div>'
        )

    return f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>{html.escape(title)} - Course Content Map</title>
<style>
:root{{--bg:#f6f7fb;--panel:#fff;--text:#172033;--muted:#667085;--border:#d9dee8;--accent:#3157d5;--accent-soft:#eef3ff;--warning:#9a3412;--warning-bg:#fff7ed}}
*{{box-sizing:border-box}}html{{scroll-behavior:smooth}}
body{{margin:0;background:var(--bg);color:var(--text);font-family:Arial,Helvetica,sans-serif;line-height:1.5}}
header{{background:var(--panel);border-bottom:1px solid var(--border);padding:26px 28px 20px}}
header h1{{margin:0 0 6px;font-size:1.8rem}}header p{{margin:0;color:var(--muted)}}
.layout{{max-width:1380px;margin:0 auto;padding:22px;display:grid;grid-template-columns:300px minmax(0,1fr);gap:22px}}
.sidebar{{position:sticky;top:16px;align-self:start;max-height:calc(100vh - 32px);overflow:auto;background:var(--panel);border:1px solid var(--border);border-radius:14px;padding:16px}}
.sidebar h2{{margin:0 0 10px;font-size:1rem}}.sidebar ol{{padding-left:1.35rem;margin:10px 0 0}}.sidebar li{{margin:.45rem 0}}
.sidebar a,.resource-link{{color:var(--accent);text-decoration:none}}.resource-link{{font-weight:600}}.resource-link:hover{{text-decoration:underline}}
.search,select{{width:100%;padding:9px 10px;border:1px solid var(--border);border-radius:8px;background:#fff;font:inherit}}
.search{{margin:8px 0 12px}}.filters{{border-top:1px solid var(--border);padding-top:12px;margin-top:10px}}
.filters-title{{font-weight:700;font-size:.9rem;margin-bottom:8px}}.filter-field{{display:block;margin:8px 0}}
.filter-field span{{display:block;color:var(--muted);font-size:.75rem;font-weight:600;margin-bottom:3px}}
.check-field{{display:block;font-size:.83rem;margin:10px 0}}.check-field input{{width:auto}}
#clear-filters{{width:100%;padding:8px 10px;border:1px solid var(--border);border-radius:8px;background:#fff;cursor:pointer}}
.results-note{{color:var(--muted);font-size:.8rem;margin-top:8px}}
.metrics{{display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:10px;margin-bottom:18px}}
.metric{{background:var(--panel);border:1px solid var(--border);border-radius:12px;padding:13px}}.metric strong{{display:block;font-size:1.35rem}}.metric span{{color:var(--muted);font-size:.84rem}}
.course-section{{background:var(--panel);border:1px solid var(--border);border-radius:15px;padding:18px;margin-bottom:18px}}
.section-heading{{display:flex;gap:12px;align-items:center}}.section-number{{width:38px;height:38px;border-radius:10px;display:flex;align-items:center;justify-content:center;background:var(--accent-soft);color:var(--accent);font-weight:700;flex:0 0 auto}}
.course-section h2{{margin:0;font-size:1.25rem}}.section-count{{color:var(--muted);font-size:.86rem;margin-top:2px}}
.section-summary{{margin:14px 0;color:var(--muted)}}.section-summary summary{{cursor:pointer;color:var(--accent);font-weight:600}}
.subgroup{{border-left:3px solid var(--border);padding-left:14px;margin:18px 0}}.subgroup h3{{margin:0 0 10px;font-size:1rem}}
.activity{{border-top:1px solid var(--border);padding:13px 0}}.activity:first-child{{border-top:0}}.activity-title-row{{display:flex;gap:9px;align-items:flex-start}}
.activity h4{{margin:1px 0 5px;font-size:.98rem;font-weight:600}}.item-heading{{min-width:0}}
.family-label{{display:block;color:var(--muted);font-size:.72rem;font-weight:600;margin-bottom:1px}}
.type-badge{{display:inline-block;background:#eef2f7;border-radius:999px;padding:2px 7px;font-size:.75rem;font-weight:700;white-space:nowrap}}
.resource-list{{margin:4px 0 0 78px;padding-left:1.15rem}}.resource-list li{{margin:5px 0}}
.item-meta{{display:block;color:var(--muted);font-size:.78rem}}.activity-note{{margin-left:78px;color:var(--muted);font-size:.85rem}}
.metadata{{margin:7px 0 0 78px;color:var(--muted);font-size:.8rem}}.metadata summary{{cursor:pointer;color:var(--accent);font-weight:600}}
.metadata dl{{margin:7px 0 0}}.metadata dl div{{display:grid;grid-template-columns:120px 1fr;gap:8px;margin:3px 0}}
.metadata dt{{font-weight:600}}.metadata dd{{margin:0}}
.unresolved{{display:inline-block;margin-left:8px;color:var(--warning);background:var(--warning-bg);border-radius:6px;padding:1px 6px;font-size:.76rem}}
.hidden-by-filter{{display:none!important}}.footer-note{{color:var(--muted);font-size:.85rem;margin:18px 0}}
@media(max-width:900px){{.layout{{grid-template-columns:1fr;padding:12px}}.sidebar{{position:static;max-height:none}}.metrics{{grid-template-columns:repeat(2,minmax(0,1fr))}}.resource-list,.activity-note,.metadata{{margin-left:0}}}}
@media(max-width:520px){{.metrics{{grid-template-columns:1fr}}}}
</style></head><body>
<header><h1>{html.escape(title)}</h1><p>Current Moodle Course Structure - generated from Moodle audit and extracted resources</p></header>
<div class="layout"><aside class="sidebar"><h2>Contents</h2>
<label for="map-search">Search this course map</label>
<input class="search" id="map-search" type="search" placeholder="Search course item, file or provider">
{filters_html}<div class="results-note" id="results-note"></div><ol>{"".join(toc)}</ol></aside>
<main><div class="metrics">
<div class="metric"><strong>{len(sections)}</strong><span>sections</span></div>
<div class="metric"><strong>{total_course_items}</strong><span>course items shown</span></div>
<div class="metric"><strong>{linked}</strong><span>clickable resource/link rows</span></div>
<div class="metric"><strong>{external}</strong><span>external link rows</span></div>
</div>{"".join(body)}
<p class="footer-note">This map represents the audited current Moodle structure. "Course item" is used as an umbrella term for Moodle resources/content and learning activities. Moodle labels are shown as subheadings and are not counted as displayed course items. Optional review metadata and filters are shown only when supported by the audit data; missing values are not inferred. The mapper does not infer a new pedagogic structure.</p>
</main></div>
<script>
(()=>{{
 const search=document.getElementById('map-search');
 const items=[...document.querySelectorAll('.activity')];
 const sections=[...document.querySelectorAll('.course-section')];
 const selects=[...document.querySelectorAll('select[data-filter]')];
 const duplicateBox=document.getElementById('duplicates-only');
 const clearButton=document.getElementById('clear-filters');
 const note=document.getElementById('results-note');
 const splitValues=value=>(value||'').split('|').filter(Boolean);
 const matchesSelect=(item,select)=>{{
   if(!select.value)return true;
   const values=splitValues(item.getAttribute('data-'+select.dataset.filter));
   return values.includes(select.value);
 }};
 const update=()=>{{
   const q=(search?.value||'').trim().toLowerCase();
   let visibleCount=0;
   items.forEach(item=>{{
     const searchMatch=!q||(item.dataset.search||'').includes(q);
     const selectMatch=selects.every(select=>matchesSelect(item,select));
     const duplicateMatch=!duplicateBox||!duplicateBox.checked||item.dataset.duplicate==='yes';
     const show=searchMatch&&selectMatch&&duplicateMatch;
     item.classList.toggle('hidden-by-filter',!show);
     if(show)visibleCount++;
   }});
   sections.forEach(section=>{{
     const visible=section.querySelectorAll('.activity:not(.hidden-by-filter)').length;
     const heading=(section.querySelector('h2')?.textContent||'').toLowerCase().includes(q);
     section.classList.toggle('hidden-by-filter',visible===0&&!heading);
   }});
   if(note)note.textContent=`${{visibleCount}} course item${{visibleCount===1?'':'s'}} matching`;
 }};
 search?.addEventListener('input',update);
 selects.forEach(select=>select.addEventListener('change',update));
 duplicateBox?.addEventListener('change',update);
 clearButton?.addEventListener('click',()=>{{
   if(search)search.value='';
   selects.forEach(select=>select.value='');
   if(duplicateBox)duplicateBox.checked=false;
   update();
 }});
 update();
}})();
</script></body></html>"""


def parse_args(argv=None):
    p = argparse.ArgumentParser(description="Create clickable HTML and editable Word maps of Moodle course items from Moodle Course Auditor outputs.")
    p.add_argument("course_dir", type=Path, help="Course input directory containing audit/ and extracted_files/.")
    p.add_argument("-o", "--output-dir", type=Path, default=None, help="Output directory. Default: output/<input-course-folder-name>")
    p.add_argument("--bundle", action="store_true", help="Copy linked extracted resources into output/resources for a portable map.")
    p.add_argument("--include-hidden", action="store_true", help="Include hidden Moodle course items and sections.")
    p.add_argument("--title", default="", help="Override detected course title.")
    p.add_argument("--version", action="version", version=f"%(prog)s {VERSION}")
    return p.parse_args(argv)


def main(argv=None) -> int:
    args = parse_args(argv)
    course_dir = args.course_dir.resolve()
    try:
        audit_dir, extracted_dir = validate_input(course_dir)
    except FileNotFoundError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    output_dir = args.output_dir.resolve() if args.output_dir else (Path.cwd()/"output"/course_dir.name).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    sections = read_csv(audit_dir/"sections.csv")
    activities = read_csv(audit_dir/"activities.csv")
    placements = read_csv(audit_dir/"content_placement_inventory.csv")
    title = clean(args.title) or discover_course_title(audit_dir, course_dir.name)
    print(f"Course: {title}")
    print(f"Input:  {course_dir}")
    print(f"Output: {output_dir}")
    print(f"Sections read:   {len(sections)}")
    print(f"Course items read (activities.csv): {len(activities)}")
    print(f"Placements read: {len(placements)}")
    resource_index = ResourceIndex(extracted_dir)
    duplicate_candidates = load_duplicate_candidates(audit_dir)
    section_models, map_rows, unresolved = map_course(
        sections,
        activities,
        placements,
        resource_index,
        output_dir,
        args.bundle,
        args.include_hidden,
        duplicate_candidates,
    )
    (output_dir/"index.html").write_text(render_html(title, section_models, map_rows), encoding="utf-8")
    try:
        write_docx(output_dir/"content_map.docx", title, section_models, map_rows)
    except RuntimeError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 3
    write_csv_output(output_dir/"content_map.csv", map_rows)
    write_unresolved_csv(output_dir/"unresolved_items.csv", unresolved)
    write_report(output_dir/"mapping_report.md", title, section_models, map_rows, unresolved, args.bundle)
    print("\nDone.")
    print(f"  Sections mapped:      {len(section_models)}")
    print('  Course items mapped:  ' + str(sum(real_course_item_count(s) for s in section_models)))
    print(f"  Mapping rows:         {len(map_rows)}")
    print(f"  Local resource links: {sum(r.link_type == 'local-file' for r in map_rows)}")
    print(f"  External links:       {sum(r.link_type == 'external' for r in map_rows)}")
    print(f"  Normalized matches:   {sum(r.mapping_status == 'normalized-filename' for r in map_rows)}")
    print(f"  Fuzzy matches:        {sum(r.mapping_status.startswith('high-confidence-fuzzy:') for r in map_rows)}")
    print(f"  Ambiguous matches:    {sum(r.mapping_status.startswith('ambiguous-') for r in map_rows)}")
    print(f"  Unresolved issues:    {len(unresolved)}")
    if duplicate_candidates:
        print(f"  Duplicate candidates: {len(duplicate_candidates)}")
    print(f"  HTML map:             {output_dir/'index.html'}")
    print(f"  Word working copy:    {output_dir/'content_map.docx'}")
    print(f"  CSV map:              {output_dir/'content_map.csv'}")
    print(f"  QA report:            {output_dir/'mapping_report.md'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
